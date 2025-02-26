from fetchers import fetch_html_with_requests, fetch_html_with_selenium, fetch_html_with_scrapy
from parsers import parse_with_beautifulsoup, parse_with_scrapy
from matching import fuzzy_match
from adversarial import perform_adversarial_attack, variable_analysis
from validation import validate_input
from config import user_input, urls
import os
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document

def save_as_pdf(data, filename):
    c = canvas.Canvas(filename, pagesize=letter)
    width, height = letter
    c.drawString(100, height - 100, "Extracted Information")
    y = height - 120

    for key, value in data.items():
        text = f"{key}: {value}"
        c.drawString(100, y, text)
        y -= 20

    c.save()

def save_as_doc(data, filename):
    doc = Document()
    doc.add_heading('Extracted Information', 0)

    for key, value in data.items():
        doc.add_paragraph(f"{key}: {value}")

    doc.save(filename)

def main():
    data_folder = "extracted_data"
    if not os.path.exists(data_folder):
        os.makedirs(data_folder)

    # Validate user input
    validated_input = validate_input(user_input)

    try:
        # Fetch and parse data
        ssa_html = fetch_html_with_requests(urls['ssa'], {"ssn": validated_input["ssn"]})
        ssa_data = parse_with_beautifulsoup(ssa_html)
        dmv_html = fetch_html_with_selenium(urls['dmv'])
        dmv_data = parse_with_scrapy(dmv_html)
        credit_bureau_html = fetch_html_with_requests(urls['credit_bureau'], {"ssn": validated_input["ssn"]})
        credit_bureau_data = parse_with_beautifulsoup(credit_bureau_html)
        public_records_html = fetch_html_with_selenium(urls['public_records'])
        public_records_data = parse_with_scrapy(public_records_html)
        third_party_html = fetch_html_with_requests(urls['third_party'], {"email": validated_input["email"]})
        third_party_data = parse_with_beautifulsoup(third_party_html)

        # Perform adversarial attack
        perform_adversarial_attack()

        # Perform variable analysis
        ssa_data = variable_analysis(ssa_data)
        dmv_data = variable_analysis(dmv_data)
        credit_bureau_data = variable_analysis(credit_bureau_data)
        public_records_data = variable_analysis(public_records_data)
        third_party_data = variable_analysis(third_party_data)

        # Fuzzy match user input with fetched data
        fuzzy_matches = {
            "ssa": fuzzy_match(user_input, ssa_data),
            "dmv": fuzzy_match(user_input, dmv_data),
            "credit_bureau": fuzzy_match(user_input, credit_bureau_data),
            "public_records": fuzzy_match(user_input, public_records_data),
            "third_party": fuzzy_match(user_input, third_party_data)
        }

        # Combine all data
        combined_data = {
            "SSA Data": ssa_data,
            "DMV Data": dmv_data,
            "Credit Bureau Data": credit_bureau_data,
            "Public Records Data": public_records_data,
            "Third-Party Data": third_party_data,
            "Fuzzy Matches": fuzzy_matches
        }

        # Check if property deeds exist
        if public_records_data.get('deed', False) == "Yes":
            combined_data["Property Deeds"] = "Yes"

        # Save combined data as a PDF
        pdf_filename = os.path.join(data_folder, "extracted_information.pdf")
        save_as_pdf(combined_data, pdf_filename)

        # Save combined data as a DOC file
        doc_filename = os.path.join(data_folder, "extracted_information.docx")
        save_as_doc(combined_data, doc_filename)

        print(f"Extracted information saved as {pdf_filename}")
        print(f"Extracted information saved as {doc_filename}")

    except Exception as e:
        print(f"An error occurred: {e}")

if __name__ == "__main__":
    main()
